import sys; sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage, ImageDraw, ImageFont
import io, os

BASE = r"c:\Users\r.gonzalez_fluxsolar.LAPTOP-FLUX-ECO\Downloads\Analisis_UCEN_v2\UNIVERSO_918\NOTEBOOKS"
OUT  = r"c:\Users\r.gonzalez_fluxsolar.LAPTOP-FLUX-ECO\Downloads\Analisis_UCEN_v2\Informe_UCEN_Reorganizado_v2.docx"

def src(name): return os.path.join(BASE, name)

doc = Document()
for s in doc.sections:
    s.top_margin    = Inches(0.85)
    s.bottom_margin = Inches(0.85)
    s.left_margin   = Inches(0.70)
    s.right_margin  = Inches(0.70)

FULL_W = Inches(7.1)   # ancho al insertar en Word (8.5 - 0.7*2 = 7.1")
ORANGE = RGBColor(0xFF, 0x6B, 0x35)
BLUE   = RGBColor(0x19, 0x76, 0xD2)
DARK   = RGBColor(0x1A, 0x23, 0x7E)
GRAY   = RGBColor(0x37, 0x47, 0x4F)

# ── Fuentes para Pillow ────────────────────────────────────────────────────────
def _pil_font(size=22):
    for path in [
        "C:/Windows/Fonts/calibrib.ttf",
        "C:/Windows/Fonts/arialbd.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibri.ttf",
    ]:
        try: return ImageFont.truetype(path, size)
        except: pass
    return ImageFont.load_default()

FONT_LBL  = _pil_font(22)
FONT_CAP  = _pil_font(18)

# ── Imagen con borde delgado negro (BytesIO) ──────────────────────────────────
def _bordered(path, border=3):
    im   = PILImage.open(path).convert("RGB")
    draw = ImageDraw.Draw(im)
    draw.rectangle([0, 0, im.width-1, im.height-1], outline=(0,0,0), width=border)
    buf = io.BytesIO()
    im.save(buf, format="PNG", dpi=(150,150))
    buf.seek(0)
    return buf

# ── Imagen compuesta: dos PNG lado a lado ────────────────────────────────────
def _composite(f1, f2, lbl1, lbl2, c1, c2,
               target_h=520, gap=28, border=3, pad=6, label_h=34):
    p1, p2 = src(f1), src(f2)
    missing = [f for f, p in [(f1,p1),(f2,p2)] if not os.path.exists(p)]
    if missing:
        print(f"  !! Falta: {missing}")
        return None

    def to_h(p, h):
        im = PILImage.open(p).convert("RGB")
        w  = int(im.width * h / im.height)
        return im.resize((w, h), PILImage.LANCZOS)

    im1 = to_h(p1, target_h)
    im2 = to_h(p2, target_h)

    # Recuadros con borde y padding
    box1_w = im1.width  + 2*(border+pad)
    box2_w = im2.width  + 2*(border+pad)
    box_h  = target_h   + 2*(border+pad)
    total_w = box1_w + gap + box2_w
    total_h = label_h + box_h

    canvas = PILImage.new("RGB", (total_w, total_h), (255,255,255))
    draw   = ImageDraw.Draw(canvas)

    def draw_box(x, im, lbl, color_rgb):
        # etiqueta centrada sobre el recuadro
        draw.text((x + (im.width + 2*(border+pad))//2, label_h//2),
                  lbl, fill=color_rgb, font=FONT_LBL, anchor="mm")
        # borde
        draw.rectangle([x, label_h,
                         x + im.width + 2*(border+pad) - 1,
                         label_h + box_h - 1],
                       outline=(0,0,0), width=border)
        # imagen pegada
        canvas.paste(im, (x + border + pad, label_h + border + pad))

    draw_box(0,           im1, lbl1, c1)
    draw_box(box1_w+gap,  im2, lbl2, c2)

    buf = io.BytesIO()
    canvas.save(buf, format="PNG", dpi=(150,150))
    buf.seek(0)
    return buf

# ── Insertar imagen en Word ───────────────────────────────────────────────────
def _insert(buf_or_path, width=None):
    if width is None: width = FULL_W
    if isinstance(buf_or_path, str):
        buf_or_path = _bordered(buf_or_path)
    doc.add_picture(buf_or_path, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── API pública ───────────────────────────────────────────────────────────────
def pair(f1, f2, lbl1="SAT Nota", lbl2="SAT % Recomendación",
         c1=None, c2=None, caption=None):
    if c1 is None: c1 = (230, 100,  40)  # naranja
    if c2 is None: c2 = ( 20, 100, 200)  # azul
    buf = _composite(f1, f2, lbl1, lbl2, c1, c2)
    if buf: _insert(buf)
    if caption: cap(caption)
    doc.add_paragraph()

def single(f, label=None, label_color=None, caption=None, width=None):
    if label:
        pl = doc.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = pl.add_run(label)
        rl.bold = True; rl.font.size = Pt(10)
        if label_color: rl.font.color.rgb = label_color
    path = src(f)
    if os.path.exists(path): _insert(path, width=width or FULL_W)
    else: doc.add_paragraph(f"[{f}]")
    if caption: cap(caption)
    doc.add_paragraph()

def stacked(f1, f2, lbl1="SAT Nota", lbl2="SAT % Recomendación",
            c1=None, c2=None, caption=None, width=None):
    if c1 is None: c1 = ORANGE
    if c2 is None: c2 = BLUE
    single(f1, label=lbl1, label_color=c1, width=width)
    single(f2, label=lbl2, label_color=c2, width=width)
    if caption: cap(caption)

# ── Texto helpers ─────────────────────────────────────────────────────────────
def cap(text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.italic = True; r.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def h1(t):
    p = doc.add_heading(t, level=1)
    if p.runs: p.runs[0].font.color.rgb = DARK

def h2(t):
    p = doc.add_heading(t, level=2)
    if p.runs: p.runs[0].font.color.rgb = GRAY

def body(t): doc.add_paragraph(t)
def bul(t):  doc.add_paragraph(t, style="List Bullet")
def sep():   doc.add_paragraph()

def nota(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    r = p.add_run(f"Nota: {t}")
    r.italic = True; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x54, 0x6E, 0x7A)


# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("ENTREGA PRODUCTO 3", 0)
h2("Análisis de Iniciativas de Formación Docente e Impacto en el Rendimiento Académico y Satisfacción Estudiantil 2022–2025")
body("Universidad Central de Chile  ·  Mayo 2026")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. UNIVERSO
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Universo de Análisis")
h2("1.1 Cascada — Evaluación SAT")
body("El universo de estudio está compuesto por 917 docentes jerarquizados. Se aplican dos filtros sucesivos para determinar los subgrupos analíticamente pertinentes:")
bul("917 docentes jerarquizados (universo total): 493 con perfil completo (nómina + dotación) y 424 solo nómina.")
bul("357 docentes (39%) participaron en al menos una instancia de formación — grupo tratamiento. Los 560 restantes (61%) conforman el grupo control.")
bul("197 docentes APTOS P3 (21%): con evaluación SAT válida en período baseline y resultado — 154 tipo TALLER, 36 tipo DIPLOMADO y 7 tipo PROYECTO.")
nota("Solo se incluyeron evaluaciones SAT con cobertura ≥ 40% (CM-1). El SAT de cada docente se promedió ponderado por número de alumnos evaluados (CM-2).")
sep()

h2("1.2 Cascada — Evaluación de Jefes (EDD)")
bul("1.413 registros EDD válidos dentro del universo 917 (486 docentes únicos).")
bul("702 registros → 235 docentes formados; 711 → 251 docentes sin formación.")
bul("Formados concentran 74% en Muy Bueno vs 62% del control — brecha de 12 pp especialmente marcada en 2025.")
sep()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. METODOLOGÍA
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Marco Metodológico")
h2("2.1 El indicador z-score")
body("El z-score expresa la posición relativa del docente dentro de su facultad y período:")
body("z = (SAT docente − promedio facultad·período) / desviación estándar facultad·período")
bul("z = 0: en el promedio exacto de su facultad ese semestre.")
bul("z = +1: supera a sus colegas en 1 desviación estándar (~84° percentil).")
bul("z = −1: está por debajo en 1 desviación estándar (~16° percentil).")

h2("2.2 Dos dimensiones de satisfacción")
bul("SAT Nota: evaluación numérica del docente (escala 1–7), z-scoreada por facultad × período.")
bul("SAT % Recomendación: porcentaje de alumnos que responden 'Sí' a '¿Recomendaría a este/a docente?', z-scoreada. Correlación con SAT Nota: r = 0.893.")
body("Ambas métricas se presentan siempre en paralelo. Su consistencia es evidencia de robustez.")
sep()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. ANÁLISIS DE IMPACTO SAT
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Análisis de Impacto en Evaluación Docente")

h2("3.1 Trayectoria Z-score: Baseline → Resultado")
pair("G1_linea_z_918.png", "G1.2_linea_zbin_918.png",
     caption="G1 / G1.2 — Trayectoria z-score por tipo de formación (n=197 aptos P3)")
bul("Los tres grupos parten sobre el promedio de su facultad (z > 0) antes de formarse — efecto de selección: quienes participan en formación tienden a ser docentes ya bien evaluados.")
bul("Taller (n=154) y Diplomado (n=36) muestran trayectorias estables en ambas métricas: el z-score post no difiere significativamente del baseline.")
bul("Proyecto (n=7) registra la caída más pronunciada en ambas dimensiones, pero el tamaño muestral hace el promedio muy sensible a casos individuales — resultado referencial, no conclusivo.")
sep(); doc.add_page_break()

h2("3.2 Delta Z-score y Porcentaje de Mejora por Tipo")
pair("G2_barras_tipo_918.png", "G2.2_barras_tipo_bin_918.png",
     caption="G2 / G2.2 — Δ z-score y % docentes que mejoran, por tipo de formación")
bul("Los tres tipos registran Δz negativo o cercano a cero en ambas métricas: la formación no produce un salto sistemático en el promedio grupal.")
bul("El % de mejora individual cuenta otra historia: ~52–56% de diplomados y ~49% de talleristas mejoró su posición relativa — la distribución es bimodal, no hay caída uniforme.")
bul("La aparente contradicción (Δz negativo con ~50% de mejora) se explica porque el promedio es sensible a quienes retroceden más. Ambos datos son verdaderos simultáneamente.")
sep(); doc.add_page_break()

h2("3.3 Impacto por Antigüedad del Docente")
pair("G3_heatmap_antiguedad_918.png", "G3.2_heatmap_antiguedad_bin_918.png",
     caption="G3 / G3.2 — Δ z-score por antigüedad × tipo de formación")
bul("Los docentes noveles (0–4 años, n=57) obtienen el mayor beneficio del Diplomado (Δz=+0.18 en nota), replicado en % recomendación — el programa es especialmente efectivo en quienes aún consolidan su práctica.")
bul("Los docentes de mayor trayectoria (15+ años) responden mejor al Taller (Δz=+0.19 en nota) — patrón consistente en ambas métricas.")
bul("El grupo de antigüedad media (5–14 años) muestra resultados cercanos a cero — la cohorte más estable pero que menos se desplaza con la intervención.")
sep(); doc.add_page_break()

h2("3.4 Impacto por Jerarquía Académica")
pair("G4_heatmap_jerarquia_918.png", "G4.2_heatmap_jerarquia_bin_918.png",
     caption="G4 / G4.2 — Δ z-score por jerarquía × tipo de formación")
bul("Instructores y Asistentes Docentes exhiben estabilidad en Taller y leve mejora en Diplomado en ambas métricas.")
bul("Asociados Docentes (n=35 en Taller) presentan Δz negativo en nota y recomendación, aunque sin diferencia estadísticamente significativa (ANOVA p=0.804).")
bul("Ninguna jerarquía mejora en una métrica y retrocede en la otra — la consistencia entre G4 y G4.2 otorga solidez interna al análisis.")
sep(); doc.add_page_break()

h2("3.5 Impacto por Sexo del Docente")
stacked("G5_barras_sexo_918.png", "G5.2_barras_sexo_bin_918.png",
        caption="G5 / G5.2 — Δ z-score por sexo × tipo de formación")
bul("Las mujeres (59% de los formados, n=116) muestran Δz positivo en Taller y Diplomado en ambas métricas; los hombres retroceden en los tres tipos.")
bul("En Taller el 51% de las mujeres mejoró en nota (53% en recomendación) vs 45% de los hombres (44%). Diferencial consistente en dos indicadores independientes.")
bul("Proyecto (F=2, M=5) no es estadísticamente interpretable; su resultado no debe generalizarse.")
sep(); doc.add_page_break()

h2("3.6 Formados vs Control: Evolución por Período")
pair("G6_control_vs_tratamiento_918.png", "G6.2_control_bin_918.png",
     caption="G6 / G6.2 — Z-score formados vs control por período 2023–2025")
bul("Los docentes formados superan al grupo control en los seis períodos sin excepción, con brecha promedio de +0.18z (nota) y +0.15z (recomendación) — equivalente a ~6–7 percentiles.")
bul("Diferencia estadísticamente significativa en los mismos tres períodos en ambas métricas: 2024-01**, 2024-02**, 2025-02*** (nota) y 2024-01**, 2024-02*, 2025-02* (recomendación).")
bul("La coincidencia entre métricas confirma que la ventaja de los formados no es artefacto de una sola pregunta.")
sep(); doc.add_page_break()

h2("3.7 Evolución por Jerarquía: Talleres")
pair("G7_taller_jerarquia_918.png", "G7.2_taller_jerarquia_bin_918.png",
     caption="G7 / G7.2 — Z-score por jerarquía: Taller (n=154 aptos P3, 2023–2025)")
bul("Asistentes Docentes (n=52) parten bajo la media y alcanzan su peak durante los talleres (2024), cerrando 2025-02 en posición favorable en ambas métricas.")
bul("Instructores Docentes (n=44) parten sobre el promedio y se mantienen estables — sin gran ganancia adicional pero sin deterioro.")
bul("Asociados Docentes (n=35) permanecen bajo la media en todos los períodos en ambas métricas — el taller no logra reposicionarlos dentro de su facultad.")
sep(); doc.add_page_break()

h2("3.8 Evolución por Jerarquía: Diplomado")
pair("G7_diplomado_jerarquia_918.png", "G7.2_diplomado_jerarquia_bin_918.png",
     caption="G7 / G7.2 — Z-score por jerarquía: Diplomado (n=36 aptos P3, 2023–2025)")
bul("Instructores Docentes (n=18) son el caso de mayor éxito: parten bajo la media, escalan durante el diplomado y consolidan en z=+0.44 en 2025-02. Trayectoria ascendente replicada en % recomendación.")
bul("Asistentes Docentes (n=12) se mantienen sobre el promedio en todos los períodos en ambas métricas.")
bul("El ANOVA no detecta diferencias significativas entre jerarquías (p=0.718), pero el patrón descriptivo para Instructores es consistente en nota y recomendación.")
sep(); doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. SAT vs NOTAS ALUMNOS
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Relación entre Evaluación Docente y Rendimiento Estudiantil")
single("G8_scatter_sat_notas_918.png",
       caption="G8 — Dispersión SAT vs Nota promedio alumnos (n=17.248 secciones, 6 períodos 2023–2025)")
bul("La correlación SAT–notas es positiva y significativa (r=0.28, p<0.001), pero modesta: explica solo el 8% de la varianza. El 92% restante responde a factores externos — dificultad, perfil del curso, carrera.")
bul("Los docentes formados (naranja) aparecen más concentrados en el sector de SAT alto, consistente con G6. El gráfico no permite atribuir causalidad directa a la formación.")
bul("Los cuadrantes de 'mal evaluado + notas bajas' están prácticamente vacíos de docentes formados: la formación no coexiste con bajo desempeño.")
sep(); doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. PRUEBAS ESTADÍSTICAS
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Validación Estadística")

h2("5.1 Prueba t — Z-score SAT Nota")
single("G11_pruebas_t_918.png",
       caption="G11 — Pruebas t SAT Nota: Pareada (Panel A) e Independiente por período (Panel C)")
bul("Prueba t Pareada (Panel A): todos los tipos ns. La formación no desplaza al docente respecto a sus colegas de facultad.")
bul("Prueba t Independiente (Panel C): significativa en 2024-01 (p=0.003, d=+0.260), 2024-02 (p=0.009, d=+0.231) y 2025-02 (p=0.001, d=+0.302).")
bul("El efecto más fuerte es 2025-02 (Cohen d=+0.302, pequeño-mediano): los formados superan al control en casi un tercio de desviación estándar.")
sep()

h2("5.2 Prueba t — Z-score % Recomendación")
single("G11.2_pruebas_t_bin_918.png",
       caption="G11.2 — Pruebas t % Recomendación: Pareada (Panel A) e Independiente por período (Panel C)")
bul("Prueba t Pareada: idéntico a SAT Nota — todos ns. La formación no mueve la tasa de recomendación del propio docente.")
bul("Prueba t Independiente: significativa en los mismos tres períodos (2024-01**, 2024-02*, 2025-02*), con efecto levemente menor (d máx.=+0.230).")
bul("La coincidencia entre métricas — mismos períodos significativos, misma dirección — es la evidencia de robustez más sólida del análisis.")
sep(); doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. EVALUACIÓN DE JEFES (EDD)
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Evaluación de Desempeño Docente (EDD) — Perspectiva de Jefes")

h2("6.1 Evolución EDD y Distribución de Concepto")
pair("G9_edd_evolucion_918.png", "G10_concepto_edd_918.png",
     lbl1="Evolución EDD (2022–2025)", lbl2="Distribución Concepto EDD",
     c1=(80,80,80), c2=(80,80,80),
     caption="G9 / G10 — EDD: evolución por año y distribución de concepto, formados vs sin formación")
bul("Los formados superan al control en EDD durante los cuatro años. Brecha más crítica en 2025: 0.864 vs 0.604 (+0.260 puntos, 26 pp). El control no logra recuperarse al ritmo de los formados.")
bul("Formados concentran 74.3% en Muy Bueno vs 65.2% del control (+9.1 pp). La categoría Deficiente casi desaparece en formados (0.5% vs 1.6%).")
bul("Diferencias estadísticamente significativas: prueba t en 2024 (p=0.002, d=+0.339) y 2025 (p<0.001, d=+0.955). Chi-cuadrado global p=0.001.")
sep()

h2("6.2 Validación Estadística EDD")
single("G12_pruebas_t_edd_918.png",
       caption="G12 — Validación estadística EDD: Prueba t por año (Panel A) y Chi-cuadrado concepto (Panel B)")
bul("Prueba t independiente: significativa en 2024 (**) y 2025 (***). Efecto global: t=+9.615, p<0.001, Cohen d=+0.522.")
bul("Chi-cuadrado: distribución de conceptos difiere entre grupos a nivel global (p=0.001, **) y en 2024 (p=0.003, **).")
bul("Cohen d=+0.955 en 2025 es el efecto más grande de todo el análisis — la dimensión más sensible al impacto de la formación.")
sep()

h2("6.3 Análisis por Jerarquía: Diplomado y Taller")
stacked("G13_ttest_jerarquia_diplomado_918.png",
        "G13.2_ttest_jerarquia_taller_918.png",
        lbl1="Diplomado — Prueba t por Jerarquía",
        lbl2="Taller — Prueba t por Jerarquía",
        c1=GRAY, c2=GRAY,
        caption="G13 / G13.2 — ANOVA + comparaciones por pares: impacto por jerarquía en Diplomado y Taller")
bul("ANOVA Diplomado: F(2,33)=0.335, p=0.718, η²=0.021 — no significativo. El impacto es estadísticamente equivalente entre jerarquías.")
bul("ANOVA Taller: F(6,144)=0.505, p=0.804, η²=0.021 — no significativo. Ninguna de las 21 comparaciones por pares alcanza significancia tras corrección Bonferroni.")
bul("La jerarquía académica no modera el impacto de la formación en SAT — el efecto es homogéneo a lo largo de toda la escala jerárquica.")
sep(); doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Conclusiones")

h2("7.1 Hallazgos principales")
bul("La formación docente alcanza al 39% de los jerarquizados. Los formados parten con posición relativa favorable y esa ventaja se mantiene y amplía — especialmente en 2024 y 2025.")
bul("El impacto sobre satisfacción estudiantil es neutral en términos pre/post del propio docente, pero los formados como grupo están consistentemente mejor posicionados que el control en 3 de 6 períodos (Cohen d hasta +0.302).")
bul("El impacto más fuerte es en EDD: Cohen d=+0.955 en 2025, efecto grande. La formación se asocia nítidamente con mejores evaluaciones de desempeño desde los jefes directos.")
bul("Ambas métricas de satisfacción (r=0.893) apuntan en la misma dirección en todos los períodos y subgrupos — robustez interna: los resultados no dependen de un solo indicador.")

h2("7.2 Matices y limitaciones")
bul("El diseño es observacional. No es posible atribuir causalidad directa: los docentes que participan pueden tener características previas que explican parte de su mejor desempeño.")
bul("El grupo Proyecto (n=7) no es estadísticamente interpretable de forma independiente. Sus resultados son ilustrativos y no generalizables.")
bul("La jerarquía académica no modera el impacto en ninguna modalidad (ANOVA p>0.70). El efecto es equivalente a lo largo de toda la escala jerárquica.")

sep()
body("Informe generado automáticamente  |  Universo 917  |  Producto 3 UCEN  |  Mayo 2026")

doc.save(OUT)
print(f"Guardado: {OUT}")
