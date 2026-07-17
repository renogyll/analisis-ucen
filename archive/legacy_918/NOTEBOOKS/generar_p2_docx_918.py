import sys; sys.stdout.reconfigure(encoding="utf-8")
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
import numpy as np
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE    = os.path.dirname(__file__)
PROC    = os.path.join(BASE, "..", "PROCESADO")
OUT_DOC = os.path.join(BASE, "..", "ENTREGABLE_P2_918.docx")

matplotlib.rcParams.update({"font.family": "DejaVu Sans", "font.size": 11})

C_TALLER    = "#1976D2"
C_DIPLOMADO = "#388E3C"
C_PROYECTO  = "#E65100"
C_UDD       = "#1565C0"
C_DTDE      = "#6A1B9A"
C_VRIIP     = "#2E7D32"
C_PART      = "#1976D2"
C_NO        = "#CFD8DC"

# ── Cargar datos ──────────────────────────────────────────────────────────────
part = pd.read_csv(os.path.join(PROC, "participacion_p2_918.csv"),
                   encoding="utf-8-sig", dtype=str)
doc917 = pd.read_csv(os.path.join(PROC, "docente_918.csv"),
                     encoding="utf-8-sig", dtype={"rut_key": str})
doc917["rut_key"] = doc917["rut_key"].str.strip()
doc917["jerarquia"] = doc917["jerarquia"].fillna(
    doc917["jerarquia_dot"].str.replace(r"^PROFESOR\s+", "", regex=True))
doc917.loc[doc917["jerarquia"] == "NO INFORMA", "jerarquia"] = None
REMAP_ANT = {"15-19":"15+","20-24":"15+","25-29":"15+","30+":"15+"}
doc917["tramo_ant"] = doc917["tramo_antiguedad"].replace(REMAP_ANT)

part["anio"] = part["anio"].astype(str).str.strip()
ruts_part = set(part["rut_key"].unique())
n_total   = len(doc917)
n_part    = len(ruts_part)
n_no_part = n_total - n_part

ORD_JER = ["INSTRUCTOR REGULAR","INSTRUCTOR DOCENTE","ASISTENTE REGULAR",
           "ASISTENTE DOCENTE","ASOCIADO REGULAR","ASOCIADO DOCENTE",
           "TITULAR REGULAR","TITULAR DOCENTE"]
ABREV_JER = {
    "INSTRUCTOR REGULAR":"Instr. Regular","INSTRUCTOR DOCENTE":"Instr. Docente",
    "ASISTENTE REGULAR":"Asist. Regular","ASISTENTE DOCENTE":"Asist. Docente",
    "ASOCIADO REGULAR":"Asoc. Regular","ASOCIADO DOCENTE":"Asoc. Docente",
    "TITULAR REGULAR":"Titular Regular","TITULAR DOCENTE":"Titular Docente",
}
ORD_ANT = ["0-4","5-9","10-14","15+"]

pngs = {}

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICO 1.1 — Participación global
# ══════════════════════════════════════════════════════════════════════════════
fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6),
                                gridspec_kw={"width_ratios":[1, 1.4]})
fig.suptitle("Participación en Perfeccionamiento Docente — Universo 917\n"
             f"{n_part} docentes formados · 448 conteos por modalidad (con doble conteo) · Períodos 2022–2025",
             fontsize=13, fontweight="bold")

# Dona
vals_dona = [n_part, n_no_part]
colors_dona = [C_PART, C_NO]
wedges, _ = ax1.pie(vals_dona, colors=colors_dona, startangle=90,
                    counterclock=False,
                    wedgeprops=dict(width=0.55, edgecolor="white", linewidth=2.5))
ax1.text(0, 0, f"{n_part}\ndocentes", ha="center", va="center",
         fontsize=13, fontweight="bold", color="#333333")
for wedge, lbl, v, p in zip(wedges,
    [f"Participaron\n{n_part} ({100*n_part/n_total:.1f}%)",
     f"No participaron\n{n_no_part} ({100*n_no_part/n_total:.1f}%)"],
    vals_dona, [100*n_part/n_total, 100*n_no_part/n_total]):
    ang = (wedge.theta2 + wedge.theta1)/2
    x2, y2 = 1.3*np.cos(np.radians(ang)), 1.3*np.sin(np.radians(ang))
    ha = "left" if x2>0 else "right"
    ax1.annotate(lbl, xy=(0.75*np.cos(np.radians(ang)), 0.75*np.sin(np.radians(ang))),
                 xytext=(x2,y2), arrowprops=dict(arrowstyle="-",color="#888",lw=1),
                 fontsize=9.5, ha=ha, va="center", fontweight="bold",
                 color=C_PART if v==n_part else "#888888")
ax1.set_xlim(-2,2)
ax1.set_title("Cobertura global", pad=10)

# Barras por modalidad
tipos = ["TALLER","DIPLOMADO","PROYECTO"]
colores_tipo = [C_TALLER, C_DIPLOMADO, C_PROYECTO]
n_tipo = [part[part["tipo_formacion"]==t]["rut_key"].nunique() for t in tipos]
x = np.arange(len(tipos))
bars = ax2.bar(x, n_tipo, color=colores_tipo, alpha=0.88, width=0.55, edgecolor="white")
for i, (n, color) in enumerate(zip(n_tipo, colores_tipo)):
    ax2.text(i, n+2, f"{n}\n({100*n/n_part:.0f}% de formados)",
             ha="center", fontsize=10, fontweight="bold", color=color)
ax2.set_xticks(x)
ax2.set_xticklabels(["Taller","Diplomado","Proyecto"], fontsize=12, fontweight="bold")
ax2.set_ylabel("N° docentes únicos")
ax2.set_title("Docentes únicos por modalidad\n"
              "(un docente puede aparecer en más de una barra)", pad=10, fontsize=10)
ax2.set_ylim(0, max(n_tipo)*1.25)
ax2.spines["top"].set_visible(False); ax2.spines["right"].set_visible(False)
plt.tight_layout()
p = os.path.join(BASE, "p2_g11_cobertura.png")
plt.savefig(p, dpi=150, bbox_inches="tight"); plt.close(); pngs["1.1"] = p

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICO 1.2 — Evolución por año
# ══════════════════════════════════════════════════════════════════════════════
fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))
fig.suptitle("Evolución de la Participación por Año y Tipo de Formación\n"
             f"Universo 917 · {n_part} participaron · Períodos 2022–2025",
             fontsize=13, fontweight="bold")

anios = ["2022","2023","2024","2025"]
for tipo, color in zip(["TALLER","DIPLOMADO","PROYECTO"],
                        [C_TALLER, C_DIPLOMADO, C_PROYECTO]):
    vals = [part[(part["tipo_formacion"]==tipo)&(part["anio"]==a)]["rut_key"].nunique()
            for a in anios]
    ax1.plot(anios, vals, marker="o", color=color, linewidth=2.5,
             markersize=9, label=tipo.capitalize())
    for i,(a,v) in enumerate(zip(anios,vals)):
        if v>0: ax1.text(i, v+2, str(v), ha="center", fontsize=9,
                         fontweight="bold", color=color)
n_por_anio_total = [part[part["anio"]==a]["rut_key"].nunique() for a in anios]
ax1.set_xticks(range(len(anios)))
ax1.set_xticklabels([f"{a}\n(n={n})" for a,n in zip(anios, n_por_anio_total)], fontsize=10)
ax1.set_ylabel("N° docentes únicos"); ax1.set_title("Docentes únicos por año y tipo")
ax1.legend(fontsize=10); ax1.set_ylim(0, None)
ax1.spines["top"].set_visible(False); ax1.spines["right"].set_visible(False)

# Barras apiladas absolutas
w = 0.55; x_idx = np.arange(len(anios)); bottom = np.zeros(len(anios))
for tipo, color in zip(["TALLER","DIPLOMADO","PROYECTO"],
                        [C_TALLER, C_DIPLOMADO, C_PROYECTO]):
    vals = [len(part[(part["tipo_formacion"]==tipo)&(part["anio"]==a)])
            for a in anios]
    ax2.bar(x_idx, vals, bottom=bottom, color=color, alpha=0.88,
            label=tipo.capitalize(), width=w, edgecolor="white")
    for i,v in enumerate(vals):
        if v>5: ax2.text(i, bottom[i]+v/2, str(v), ha="center", va="center",
                         fontsize=9, fontweight="bold", color="white")
    bottom += np.array(vals)
inst_por_anio = [len(part[part["anio"]==a]) for a in anios]
ax2.set_xticks(x_idx); ax2.set_xticklabels([f"{a}\n(n={n})" for a,n in zip(anios, inst_por_anio)], fontsize=10)
n_inst_total = len(part)
ax2.set_ylabel("N° instancias"); ax2.set_title(f"Total instancias por año — {n_inst_total} actividades de formación")
ax2.legend(fontsize=10)
ax2.spines["top"].set_visible(False); ax2.spines["right"].set_visible(False)
plt.tight_layout()
p = os.path.join(BASE, "p2_g12_evolucion.png")
plt.savefig(p, dpi=150, bbox_inches="tight"); plt.close(); pngs["1.2"] = p

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICO 2.1 — Participación × Jerarquía
# ══════════════════════════════════════════════════════════════════════════════
jer_validos = [j for j in ORD_JER if j in doc917["jerarquia"].values]
n_jer_total = [len(doc917[doc917["jerarquia"]==j]) for j in jer_validos]
n_jer_part  = [len(set(part[part["jerarquia"]==j]["rut_key"])) for j in jer_validos]
n_jer_no    = [t-p for t,p in zip(n_jer_total, n_jer_part)]
pct_part    = [100*p/t if t>0 else 0 for p,t in zip(n_jer_part, n_jer_total)]
labels_jer  = [f"{ABREV_JER[j]}\n(n={t})" for j,t in zip(jer_validos, n_jer_total)]

fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))
fig.suptitle("Participación en Perfeccionamiento por Jerarquía Académica\n"
             "Universo 917 · Períodos 2022–2025", fontsize=13, fontweight="bold")

x = np.arange(len(jer_validos))
ax1.bar(x, [100*p/t for p,t in zip(n_jer_part,n_jer_total)],
        color=C_PART, alpha=0.85, width=0.6, label="Participaron", edgecolor="white")
ax1.bar(x, [100*n/t for n,t in zip(n_jer_no,n_jer_total)],
        bottom=[100*p/t for p,t in zip(n_jer_part,n_jer_total)],
        color=C_NO, alpha=0.85, width=0.6, label="No participaron", edgecolor="white")
for i,(pp,t) in enumerate(zip(pct_part, n_jer_total)):
    if pp>=8: ax1.text(i, pp/2, f"{pp:.0f}%", ha="center", va="center",
                        fontsize=9, fontweight="bold", color="white")
ax1.set_xticks(x); ax1.set_xticklabels(labels_jer, fontsize=8.5)
ax1.set_ylabel("% de docentes"); ax1.set_title("% participación por jerarquía")
ax1.set_ylim(0,108); ax1.legend(fontsize=10)
ax1.spines["top"].set_visible(False); ax1.spines["right"].set_visible(False)

ax2.barh(x, pct_part, color=C_PART, alpha=0.85, height=0.6, edgecolor="white")
for i,p in enumerate(pct_part):
    ax2.text(p+0.5, i, f"{p:.0f}%  (n={n_jer_part[i]})",
             va="center", fontsize=9.5, fontweight="bold", color=C_PART)
ax2.set_yticks(x); ax2.set_yticklabels(labels_jer, fontsize=8.5)
ax2.set_xlabel("% que participó"); ax2.set_title("Ranking de participación")
ax2.set_xlim(0,100); ax2.axvline(100*n_part/n_total, color="#C62828",
             linewidth=1.5, linestyle="--", alpha=0.7,
             label=f"Media global ({100*n_part/n_total:.0f}%)")
ax2.legend(fontsize=9)
ax2.spines["top"].set_visible(False); ax2.spines["right"].set_visible(False)
plt.tight_layout()
p = os.path.join(BASE, "p2_g21_jerarquia.png")
plt.savefig(p, dpi=150, bbox_inches="tight"); plt.close(); pngs["2.1"] = p

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICO 2.2 — Participación × Facultad
# ══════════════════════════════════════════════════════════════════════════════
# Solo docentes con perfil completo (tienen unidad_facultad)
doc_fac = doc917[doc917["unidad_facultad"].notna()].copy()

FAC_ABREV = {
    "FAC. DE INGENIERÍA Y ARQUITECTURA":            "Ing. y Arquitectura",
    "FAC. ECONOMÍA, GOBIERNO Y COMUNICACIONES":     "Econ. Gob. y Com.",
    "FAC. DE EDUCACIÓN":                            "Educación",
    "FAC. DERECHO Y HUMANIDADES":                   "Derecho y Hum.",
    "FAC. DE MEDICINA Y CIENCIAS DE LA SALUD":      "Medicina y Salud",
    "CARRERAS TÉCNICAS":                            "Carreras Técnicas",
}
def abrev_fac(f):
    f = str(f).upper()
    for k,v in FAC_ABREV.items():
        if any(p in f for p in k.split(".")):
            return v
    return f[:25]

doc_fac["fac_abrev"] = doc_fac["unidad_facultad"].apply(abrev_fac)
facs = doc_fac["fac_abrev"].value_counts()
facs_ord = facs.index.tolist()

n_fac_total = [len(doc_fac[doc_fac["fac_abrev"]==f]) for f in facs_ord]
n_fac_part  = []
for f in facs_ord:
    ruts_f = set(doc_fac[doc_fac["fac_abrev"]==f]["rut_key"])
    n_fac_part.append(len(ruts_f & ruts_part))
pct_fac = [100*p/t if t>0 else 0 for p,t in zip(n_fac_part,n_fac_total)]
orden = sorted(range(len(facs_ord)), key=lambda i: pct_fac[i])
facs_ord2  = [facs_ord[i] for i in orden]
n_tot2     = [n_fac_total[i] for i in orden]
n_part2    = [n_fac_part[i] for i in orden]
pct2       = [pct_fac[i] for i in orden]

fig, ax = plt.subplots(figsize=(12, 6))
fig.suptitle("Participación en Perfeccionamiento por Facultad\n"
             "Universo 917 · Solo docentes con perfil completo (n=493)",
             fontsize=13, fontweight="bold")
y = np.arange(len(facs_ord2))
ax.barh(y, pct2, color=C_PART, alpha=0.85, height=0.6, edgecolor="white")
for i,(p,nt,np_) in enumerate(zip(pct2,n_tot2,n_part2)):
    ax.text(p+0.5, i, f"{p:.0f}%  ({np_}/{nt})",
            va="center", fontsize=10, fontweight="bold", color=C_PART)
ax.set_yticks(y); ax.set_yticklabels(facs_ord2, fontsize=10, fontweight="bold")
ax.set_xlabel("% de docentes que participaron")
ax.set_title("% participación por facultad (ordenado ascendente)", pad=10)
ax.set_xlim(0,100)
ax.axvline(100*n_part/n_total, color="#C62828", linewidth=1.5,
           linestyle="--", alpha=0.7, label=f"Media global ({100*n_part/n_total:.0f}%)")
ax.legend(fontsize=9); ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
plt.tight_layout()
p = os.path.join(BASE, "p2_g22_facultad.png")
plt.savefig(p, dpi=150, bbox_inches="tight"); plt.close(); pngs["2.2"] = p

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICO 2.3 — Participación × Antigüedad
# ══════════════════════════════════════════════════════════════════════════════
doc917_ant = doc917[doc917["tramo_ant"].notna()].copy()
n_ant_total = [len(doc917_ant[doc917_ant["tramo_ant"]==a]) for a in ORD_ANT]
n_ant_part  = [len(set(part[part["tramo_ant"]==a]["rut_key"])) for a in ORD_ANT]
n_ant_no    = [t-p for t,p in zip(n_ant_total,n_ant_part)]
pct_ant     = [100*p/t if t>0 else 0 for p,t in zip(n_ant_part,n_ant_total)]

fig, (ax1,ax2) = plt.subplots(1,2,figsize=(14,6))
n_con_ant = sum(n_ant_total)
fig.suptitle("Participación en Perfeccionamiento por Antigüedad\n"
             f"Universo 917 · {n_con_ant} con dato de antigüedad (perfil completo) · 426 sin dato (solo nómina)",
             fontsize=13, fontweight="bold")
x = np.arange(len(ORD_ANT)); w=0.35
ax1.bar(x-w/2, n_ant_part, width=w, color=C_PART, alpha=0.85,
        label="Participaron", edgecolor="white")
ax1.bar(x+w/2, n_ant_no,   width=w, color=C_NO,  alpha=0.85,
        label="No participaron", edgecolor="white")
for i,(p,n) in enumerate(zip(n_ant_part,n_ant_no)):
    ax1.text(i-w/2, p+1, str(p), ha="center", fontsize=9.5,
             fontweight="bold", color=C_PART)
    ax1.text(i+w/2, n+1, str(n), ha="center", fontsize=9.5,
             fontweight="bold", color="#888888")
ax1.set_xticks(x)
ax1.set_xticklabels([f"{a} años\n(n={t})" for a,t in zip(ORD_ANT,n_ant_total)], fontsize=10)
ax1.set_ylabel("N° docentes"); ax1.set_title("Absolutos por tramo")
ax1.legend(fontsize=10)
ax1.spines["top"].set_visible(False); ax1.spines["right"].set_visible(False)

ax2.bar(x, pct_ant, color=C_PART, alpha=0.85, width=0.55, edgecolor="white")
for i,p in enumerate(pct_ant):
    ax2.text(i, p+0.5, f"{p:.0f}%", ha="center", fontsize=11,
             fontweight="bold", color=C_PART)
ax2.axhline(100*n_part/n_total, color="#C62828", linewidth=1.5,
            linestyle="--", label=f"Media global ({100*n_part/n_total:.0f}%)")
ax2.set_xticks(x); ax2.set_xticklabels([f"{a} años" for a in ORD_ANT], fontsize=11)
ax2.set_ylabel("% que participó"); ax2.set_title("% participación por tramo")
ax2.set_ylim(0,80); ax2.legend(fontsize=9)
ax2.spines["top"].set_visible(False); ax2.spines["right"].set_visible(False)
plt.tight_layout()
p = os.path.join(BASE, "p2_g23_antiguedad.png")
plt.savefig(p, dpi=150, bbox_inches="tight"); plt.close(); pngs["2.3"] = p

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICO 3.1 — Distribución por modalidad
# ══════════════════════════════════════════════════════════════════════════════
tipos = ["TALLER","DIPLOMADO","PROYECTO"]
cols  = [C_TALLER, C_DIPLOMADO, C_PROYECTO]
n_inst_tipo = [len(part[part["tipo_formacion"]==t]) for t in tipos]
n_doc_tipo  = [part[part["tipo_formacion"]==t]["rut_key"].nunique() for t in tipos]
n_tot_inst  = sum(n_inst_tipo)

fig, (ax1,ax2) = plt.subplots(1,2,figsize=(14,6),
                               gridspec_kw={"width_ratios":[1,1.3]})
fig.suptitle("Distribución de Instancias por Modalidad de Formación\n"
             "Universo 917 · 615 instancias en universo jerarquizado",
             fontsize=13, fontweight="bold")
wedges,_ = ax1.pie(n_inst_tipo, colors=cols, startangle=90, counterclock=False,
                   wedgeprops=dict(width=0.55, edgecolor="white", linewidth=2.5))
ax1.text(0,0,f"{n_tot_inst}\ninstancias",ha="center",va="center",
         fontsize=12,fontweight="bold",color="#333")
for wedge,tipo,n,color in zip(wedges,["Taller","Diplomado","Proyecto"],
                                n_inst_tipo,cols):
    ang=(wedge.theta2+wedge.theta1)/2
    x2,y2=1.3*np.cos(np.radians(ang)),1.3*np.sin(np.radians(ang))
    ha="left" if x2>0 else "right"
    ax1.annotate(f"{tipo}\n{n} inst. ({100*n/n_tot_inst:.0f}%)",
                 xy=(0.75*np.cos(np.radians(ang)),0.75*np.sin(np.radians(ang))),
                 xytext=(x2,y2),arrowprops=dict(arrowstyle="-",color="#888",lw=1),
                 fontsize=9.5,ha=ha,va="center",fontweight="bold",color=color)
ax1.set_xlim(-2,2); ax1.set_title("Instancias por modalidad")

x=np.arange(len(tipos)); w=0.35
ax2.bar(x-w/2,n_inst_tipo,width=w,color=cols,alpha=0.85,
        label="Instancias",edgecolor="white")
ax2.bar(x+w/2,n_doc_tipo,width=w,color=cols,alpha=0.5,
        label="Docentes únicos",edgecolor="white",hatch="//")
for i,(ni,nd,c) in enumerate(zip(n_inst_tipo,n_doc_tipo,cols)):
    ax2.text(i-w/2,ni+1,str(ni),ha="center",fontsize=9.5,fontweight="bold",color=c)
    ax2.text(i+w/2,nd+1,str(nd),ha="center",fontsize=9.5,fontweight="bold",color=c)
ax2.set_xticks(x)
ax2.set_xticklabels(["Taller","Diplomado","Proyecto"],fontsize=12,fontweight="bold")
ax2.set_ylabel("N°"); ax2.set_title("Instancias vs Docentes únicos por modalidad")
ax2.legend(fontsize=10)
ax2.spines["top"].set_visible(False); ax2.spines["right"].set_visible(False)
plt.tight_layout()
p = os.path.join(BASE, "p2_g31_modalidad.png")
plt.savefig(p, dpi=150, bbox_inches="tight"); plt.close(); pngs["3.1"] = p

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICO 3.2 — Tipo × Jerarquía
# ══════════════════════════════════════════════════════════════════════════════
jer_validos2 = [j for j in ORD_JER if j in part["jerarquia"].values]
fig, ax = plt.subplots(figsize=(14,6))
fig.suptitle("Tipo de Formación por Jerarquía Académica\n"
             "Universo 917 · 357 docentes formados",
             fontsize=13, fontweight="bold")
x=np.arange(len(jer_validos2)); bottom=np.zeros(len(jer_validos2))
for tipo,color in zip(["TALLER","DIPLOMADO","PROYECTO"],[C_TALLER,C_DIPLOMADO,C_PROYECTO]):
    vals=[part[(part["tipo_formacion"]==tipo)&(part["jerarquia"]==j)]["rut_key"].nunique()
          for j in jer_validos2]
    ns_jer=[part[part["jerarquia"]==j]["rut_key"].nunique() for j in jer_validos2]
    pcts=[100*v/n if n>0 else 0 for v,n in zip(vals,ns_jer)]
    bars=ax.bar(x,pcts,bottom=bottom,color=color,alpha=0.88,
                label=tipo.capitalize(),edgecolor="white",width=0.6)
    for i,(p,v) in enumerate(zip(pcts,vals)):
        if p>=8: ax.text(i,bottom[i]+p/2,f"{p:.0f}%",ha="center",va="center",
                          fontsize=8.5,fontweight="bold",color="white")
    bottom+=np.array(pcts)
ns_jer_total=[part[part["jerarquia"]==j]["rut_key"].nunique() for j in jer_validos2]
labels2=[f"{ABREV_JER[j]}\n(n={n})" for j,n in zip(jer_validos2,ns_jer_total)]
ax.set_xticks(x); ax.set_xticklabels(labels2,fontsize=9)
ax.set_ylabel("% dentro de docentes formados de esa jerarquía")
ax.set_title("Composición del tipo de formación por jerarquía (% apilado)")
ax.set_ylim(0,108); ax.legend(fontsize=10,loc="upper right")
ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
plt.tight_layout()
p = os.path.join(BASE, "p2_g32_tipo_jerarquia.png")
plt.savefig(p, dpi=150, bbox_inches="tight"); plt.close(); pngs["3.2"] = p

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICO 4.1 — Perfil no participantes
# ══════════════════════════════════════════════════════════════════════════════
doc917["participo"] = doc917["rut_key"].isin(ruts_part)
fig, (ax1,ax2) = plt.subplots(1,2,figsize=(16,6))
fig.suptitle("Perfil de Docentes que No Han Participado en Perfeccionamiento\n"
             f"Universo 917 · {n_no_part} sin ninguna instancia registrada (2022–2025)",
             fontsize=13, fontweight="bold")
# Por jerarquía
n_jer_no2=[len(doc917[(doc917["jerarquia"]==j)&(~doc917["participo"])]) for j in jer_validos]
n_jer_si2=[len(doc917[(doc917["jerarquia"]==j)&(doc917["participo"])]) for j in jer_validos]
x=np.arange(len(jer_validos)); w=0.35
ax1.bar(x-w/2,n_jer_si2,width=w,color=C_PART,alpha=0.85,label="Participaron",edgecolor="white")
ax1.bar(x+w/2,n_jer_no2,width=w,color="#C62828",alpha=0.85,label="No participaron",edgecolor="white")
for i,(p,n) in enumerate(zip(n_jer_si2,n_jer_no2)):
    ax1.text(i-w/2,p+0.5,str(p),ha="center",fontsize=8,fontweight="bold",color=C_PART)
    ax1.text(i+w/2,n+0.5,str(n),ha="center",fontsize=8,fontweight="bold",color="#C62828")
ax1.set_xticks(x)
ax1.set_xticklabels([ABREV_JER[j] for j in jer_validos],fontsize=8.5,rotation=20,ha="right")
ax1.set_ylabel("N° docentes"); ax1.set_title("Por jerarquía: formados vs no formados")
ax1.legend(fontsize=10)
ax1.spines["top"].set_visible(False); ax1.spines["right"].set_visible(False)
# Por antigüedad
n_ant_no2=[len(doc917[(doc917["tramo_ant"]==a)&(~doc917["participo"])]) for a in ORD_ANT]
n_ant_si2=[len(doc917[(doc917["tramo_ant"]==a)&(doc917["participo"])]) for a in ORD_ANT]
x=np.arange(len(ORD_ANT))
ax2.bar(x-w/2,n_ant_si2,width=w,color=C_PART,alpha=0.85,label="Participaron",edgecolor="white")
ax2.bar(x+w/2,n_ant_no2,width=w,color="#C62828",alpha=0.85,label="No participaron",edgecolor="white")
for i,(p,n) in enumerate(zip(n_ant_si2,n_ant_no2)):
    ax2.text(i-w/2,p+0.5,str(p),ha="center",fontsize=9,fontweight="bold",color=C_PART)
    ax2.text(i+w/2,n+0.5,str(n),ha="center",fontsize=9,fontweight="bold",color="#C62828")
ax2.set_xticks(x)
ax2.set_xticklabels([f"{a} años" for a in ORD_ANT],fontsize=11)
ax2.set_ylabel("N° docentes"); ax2.set_title("Por antigüedad: formados vs no formados")
ax2.legend(fontsize=10)
ax2.spines["top"].set_visible(False); ax2.spines["right"].set_visible(False)
plt.tight_layout()
p = os.path.join(BASE, "p2_g41_brechas.png")
plt.savefig(p, dpi=150, bbox_inches="tight"); plt.close(); pngs["4.1"] = p

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICO 4.2 — Intensidad de participación
# ══════════════════════════════════════════════════════════════════════════════
intensidad = part.groupby("rut_key").size().reset_index(name="n_instancias")
bins = intensidad["n_instancias"].value_counts().sort_index()
tramos = {1:0, 2:0, 3:0, 4:0}
for v,n in bins.items():
    k = min(v,4)
    tramos[k] = tramos.get(k,0) + n
labels_int = ["1 instancia","2 instancias","3 instancias","4 o más"]
vals_int   = [tramos[1],tramos[2],tramos[3],tramos[4]]
pct_int    = [100*v/n_part for v in vals_int]
acum       = np.cumsum(pct_int)

fig, (ax1,ax2) = plt.subplots(1,2,figsize=(14,6))
fig.suptitle("Intensidad de Participación en Perfeccionamiento\n"
             f"Universo 917 · {n_part} docentes con al menos 1 instancia",
             fontsize=13, fontweight="bold")
cols_int=["#1565C0","#1976D2","#42A5F5","#90CAF9"]
bars=ax1.bar(range(4),vals_int,color=cols_int,alpha=0.88,edgecolor="white",width=0.6)
for i,(v,p) in enumerate(zip(vals_int,pct_int)):
    ax1.text(i,v+0.5,f"{v}\n({p:.0f}%)",ha="center",fontsize=10,
             fontweight="bold",color=cols_int[i])
ax1.set_xticks(range(4)); ax1.set_xticklabels(labels_int,fontsize=10)
ax1.set_ylabel("N° docentes")
ax1.set_title("Distribución por número de instancias")
ax1.spines["top"].set_visible(False); ax1.spines["right"].set_visible(False)

ax2.plot(range(4),acum,marker="o",color="#1565C0",linewidth=2.5,markersize=10)
ax2.fill_between(range(4),acum,alpha=0.15,color="#1565C0")
for i,(p,a) in enumerate(zip(pct_int,acum)):
    ax2.text(i,a+1.5,f"{a:.0f}%",ha="center",fontsize=10,
             fontweight="bold",color="#1565C0")
ax2.set_xticks(range(4)); ax2.set_xticklabels(labels_int,fontsize=10)
ax2.set_ylabel("% acumulado de docentes formados")
ax2.set_title("Curva acumulada de participación")
ax2.set_ylim(0,110)
ax2.spines["top"].set_visible(False); ax2.spines["right"].set_visible(False)
plt.tight_layout()
p = os.path.join(BASE, "p2_g42_intensidad.png")
plt.savefig(p, dpi=150, bbox_inches="tight"); plt.close(); pngs["4.2"] = p

print(f"Gráficos generados: {len(pngs)}")

# ══════════════════════════════════════════════════════════════════════════════
# ENSAMBLE WORD
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

def titulo(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
    doc.add_paragraph()

def intro(doc, texto):
    p = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.size = Pt(11)
    doc.add_paragraph()

def cascada(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33,0x33,0x33)
    doc.add_paragraph()

def grafico(doc, png_path, ancho_cm=15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(png_path, width=Cm(ancho_cm))
    doc.add_paragraph()

def bajada(doc, texto):
    p = doc.add_paragraph(style="List Bullet")
    p.runs[0].text if p.runs else p.add_run(texto)
    if not p.runs:
        p.add_run(texto)
    else:
        p.runs[0].text = texto
    p.runs[0].font.size = Pt(11)

# ── Portada ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PRODUCTO 2")
r.bold = True; r.font.size = Pt(16)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Análisis de la Participación en el Perfeccionamiento Docente")
r.bold = True; r.font.size = Pt(14)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Universo 917 docentes jerarquizados · Períodos 2022–2025")
r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x55,0x55,0x55)
doc.add_page_break()

# ── NOTA METODOLÓGICA ─────────────────────────────────────────────────────────
titulo(doc, "Nota Metodológica")
intro(doc,
    "El presente análisis cubre los 917 docentes jerarquizados de UCEN con registro en "
    "el período 2022–2025. Las instancias de formación se clasifican en tres modalidades: "
    "DIPLOMADO, TALLER y PROYECTO, provenientes de 8 archivos fuente consolidados. "
    "Las unidades ofertantes identificadas son UDD (Unidad de Desarrollo Docente), "
    "DTDE (Dirección de Transformación Digital Educativa) y VRIIP (Vicerrectoría de "
    "Investigación, Innovación y Postgrado). No se identificaron instancias de VcM ni OFEM "
    "en los registros disponibles. Los 493 docentes con perfil completo (dotación) permiten "
    "el análisis por facultad y antigüedad; los 424 restantes (solo nómina) se incluyen "
    "en la cobertura global pero no en los cruces demográficos.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 1
# ══════════════════════════════════════════════════════════════════════════════
titulo(doc, "SECCIÓN 1 — Cobertura General de Participación")

# 1.1
titulo(doc, "1.1  Participación Global")
intro(doc,
    f"Casi 4 de cada 10 docentes jerarquizados de UCEN han participado en al menos una "
    f"instancia de perfeccionamiento entre 2022 y 2025 ({n_part} de {n_total}, {100*n_part/n_total:.1f}%). "
    f"El Taller es la modalidad con mayor número de docentes únicos ({n_tipo[0]}), seguido del "
    f"Diplomado ({n_tipo[1]}) y el Proyecto ({n_tipo[2]}), evidenciando una oferta formativa "
    f"diversa pero con amplio margen de expansión en la cobertura institucional.")
cascada(doc,
    f"917 docentes jerarquizados (Universo 917)\n"
    f"  ├── {n_part} participaron al menos 1 vez ({100*n_part/n_total:.1f}%)\n"
    f"  │     ├── {n_tipo[0]:3d} con participación en TALLER\n"
    f"  │     ├── {n_tipo[1]:3d} con participación en DIPLOMADO\n"
    f"  │     └──  {n_tipo[2]:2d} con participación en PROYECTO\n"
    f"  └── {n_no_part} nunca participaron ({100*n_no_part/n_total:.1f}%)")
grafico(doc, pngs["1.1"])
bajada(doc,
    f"El {100*n_part/n_total:.1f}% del cuerpo docente jerarquizado ha participado en al menos "
    f"una instancia de formación en el período analizado. El Taller alcanza a {n_tipo[0]} docentes "
    f"únicos, siendo la modalidad de mayor cobertura, seguida del Diplomado con {n_tipo[1]} "
    f"participantes — instancias de mayor duración y profundidad formativa.")
bajada(doc,
    f"El {100*n_no_part/n_total:.1f}% restante ({n_no_part} docentes) no registra ninguna participación "
    f"en el período 2022–2025. Esta brecha de cobertura representa la principal oportunidad de "
    f"mejora en la política institucional de desarrollo académico, y su caracterización se "
    f"desarrolla en la Sección 4.")
doc.add_page_break()

# 1.2
titulo(doc, "1.2  Evolución de la Participación 2022–2025")
n_por_anio = [part[part["anio"]==a]["rut_key"].nunique() for a in anios]
anio_pico  = anios[n_por_anio.index(max(n_por_anio))]
intro(doc,
    f"La participación en instancias de perfeccionamiento muestra una tendencia creciente "
    f"entre 2022 y 2024, alcanzando su punto máximo en {anio_pico} con {max(n_por_anio)} "
    f"docentes únicos activos. El año 2025 registra menor volumen al ser un período en curso "
    f"al momento del análisis.")
cascada(doc,
    "917 docentes jerarquizados · Docentes únicos por año\n" +
    "\n".join([f"  ├── {a}: {n:3d} docentes únicos · {len(part[part['anio']==a]):3d} instancias"
               for a,n in zip(anios,n_por_anio)]))
grafico(doc, pngs["1.2"])
bajada(doc,
    f"La participación creció de {n_por_anio[0]} docentes en 2022 a {n_por_anio[2]} en 2024 "
    f"— un aumento de {n_por_anio[2]-n_por_anio[0]} docentes en dos años. El Taller lidera "
    f"el crecimiento en todos los períodos, lo que refleja la expansión de la oferta de "
    f"perfeccionamiento breve y de alta cobertura.")
bajada(doc,
    f"El año 2025 muestra {n_por_anio[3]} docentes únicos, cifra parcial dado que el análisis "
    f"se realiza durante el transcurso del año. La tendencia acumulada sugiere que la cobertura "
    f"continuará creciendo si se mantiene el ritmo de oferta formativa observado en 2024.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 2
# ══════════════════════════════════════════════════════════════════════════════
titulo(doc, "SECCIÓN 2 — Perfil de los Académicos que Participan")

# 2.1
titulo(doc, "2.1  Participación por Jerarquía Académica")
mejor_jer_idx = pct_part.index(max(pct_part))
peor_jer_idx  = pct_part.index(min(pct_part))
intro(doc,
    f"La jerarquía académica incide en el acceso al perfeccionamiento: "
    f"{ABREV_JER[jer_validos[mejor_jer_idx]]} registra la mayor tasa de participación "
    f"({pct_part[mejor_jer_idx]:.0f}%), mientras "
    f"{ABREV_JER[jer_validos[peor_jer_idx]]} presenta la menor "
    f"({pct_part[peor_jer_idx]:.0f}%). La brecha entre ambos extremos evidencia "
    f"diferencias sistemáticas en el acceso según el nivel del escalafón.")
casc = "917 docentes jerarquizados · Participación por jerarquía\n"
for j,t,p,pp in zip(jer_validos,n_jer_total,n_jer_part,pct_part):
    casc += f"  ├── {ABREV_JER[j]:18}: n={t:3d} | participaron={p:3d} ({pp:.0f}%)\n"
cascada(doc, casc)
grafico(doc, pngs["2.1"])
bajada(doc,
    f"{ABREV_JER[jer_validos[mejor_jer_idx]]} lidera la participación con un {pct_part[mejor_jer_idx]:.0f}% "
    f"de su cuerpo formado, mientras los rangos intermedios muestran tasas variables que "
    f"no siguen un patrón lineal respecto al escalafón — lo que sugiere que el acceso "
    f"al perfeccionamiento responde más a la oferta disponible que a la jerarquía per se.")
bajada(doc,
    f"La mayor brecha absoluta se concentra en {ABREV_JER[jer_validos[peor_jer_idx]]}, "
    f"con solo el {pct_part[peor_jer_idx]:.0f}% de participación (n={n_jer_part[peor_jer_idx]} de {n_jer_total[peor_jer_idx]}). "
    f"Este grupo representa una oportunidad prioritaria para ampliar la cobertura institucional "
    f"del Plan de Formación, Innovación e Investigación en Docencia.")
doc.add_page_break()

# 2.2
titulo(doc, "2.2  Participación por Facultad")
intro(doc,
    f"La cobertura del perfeccionamiento varía entre facultades, con diferencias de hasta "
    f"{max(pct_fac)-min(pct_fac):.0f} puntos porcentuales entre la más activa y la menos activa. "
    f"Este análisis considera los {len(doc_fac)} docentes con perfil completo en dotación, "
    f"excluyendo los 424 sin datos de facultad.")
casc = f"493 docentes con perfil completo · Participación por facultad\n"
for f,t,p,pp in sorted(zip(facs_ord,n_fac_total,n_fac_part,pct_fac),key=lambda x:-x[3]):
    casc += f"  ├── {f:30}: n={t:3d} | participaron={p:3d} ({pp:.0f}%)\n"
cascada(doc, casc)
grafico(doc, pngs["2.2"])
bajada(doc,
    f"La facultad con mayor cobertura alcanza el {max(pct_fac):.0f}% de participación, "
    f"mientras la de menor cobertura llega solo al {min(pct_fac):.0f}%. Esta dispersión "
    f"sugiere que el acceso al perfeccionamiento no es homogéneo entre unidades académicas "
    f"y podría estar influenciado por factores como disponibilidad horaria, cultura "
    f"departamental o pertinencia percibida de la oferta formativa.")
bajada(doc,
    f"La línea de media global ({100*n_part/n_total:.0f}%) permite identificar las facultades "
    f"que están por sobre o bajo el promedio institucional. Las unidades con menor participación "
    f"representan focos prioritarios para estrategias diferenciadas de convocatoria y "
    f"diseño de oferta formativa.")
doc.add_page_break()

# 2.3
titulo(doc, "2.3  Participación por Antigüedad")
mejor_ant_idx = pct_ant.index(max(pct_ant))
intro(doc,
    f"Los docentes con {ORD_ANT[mejor_ant_idx]} años de antigüedad muestran la mayor tasa "
    f"de participación ({pct_ant[mejor_ant_idx]:.0f}%), sugiriendo que el perfeccionamiento "
    f"se concentra en una etapa específica de la trayectoria institucional. Los docentes "
    f"de menor antigüedad presentan mayor heterogeneidad en el acceso.")
casc = "917 docentes · Participación por tramo de antigüedad\n"
for a,t,p,pp in zip(ORD_ANT,n_ant_total,n_ant_part,pct_ant):
    casc += f"  ├── {a:5} años: n={t:3d} | participaron={p:3d} ({pp:.0f}%)\n"
cascada(doc, casc)
grafico(doc, pngs["2.3"])
bajada(doc,
    f"El tramo {ORD_ANT[mejor_ant_idx]} años concentra la mayor proporción de docentes "
    f"formados ({pct_ant[mejor_ant_idx]:.0f}%), lo que puede reflejar una etapa de "
    f"consolidación profesional en la que el perfeccionamiento es más valorado o accesible. "
    f"Los docentes más nuevos (0-4 años) muestran una tasa de {pct_ant[0]:.0f}%, "
    f"posiblemente limitada por la curva de adaptación institucional.")
bajada(doc,
    f"La participación no sigue una tendencia lineal con la antigüedad — los tramos "
    f"intermedios no siempre superan a los extremos. Esto refuerza la necesidad de "
    f"estrategias diferenciadas por etapa de carrera, más que enfoques universales "
    f"de convocatoria.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 3
# ══════════════════════════════════════════════════════════════════════════════
titulo(doc, "SECCIÓN 3 — Tipos de Instancias de Formación")

# 3.1
titulo(doc, "3.1  Distribución por Modalidad")
intro(doc,
    f"El Taller es la modalidad de perfeccionamiento más extendida con {n_inst_tipo[0]} instancias "
    f"({100*n_inst_tipo[0]/n_tot_inst:.0f}% del total), seguido del Diplomado "
    f"({n_inst_tipo[1]} instancias, {100*n_inst_tipo[1]/n_tot_inst:.0f}%) y el Proyecto "
    f"({n_inst_tipo[2]} instancias, {100*n_inst_tipo[2]/n_tot_inst:.0f}%). Cada modalidad "
    f"responde a distintas profundidades formativas y niveles de compromiso docente.")
cascada(doc,
    f"615 instancias en universo 917 · Por modalidad\n"
    f"  ├── TALLER    : {n_inst_tipo[0]:3d} instancias · {n_doc_tipo[0]:3d} docentes únicos\n"
    f"  ├── DIPLOMADO : {n_inst_tipo[1]:3d} instancias · {n_doc_tipo[1]:3d} docentes únicos\n"
    f"  └── PROYECTO  :  {n_inst_tipo[2]:2d} instancias ·  {n_doc_tipo[2]:2d} docentes únicos")
grafico(doc, pngs["3.1"])
bajada(doc,
    f"El Taller concentra el {100*n_inst_tipo[0]/n_tot_inst:.0f}% de las instancias y alcanza "
    f"a {n_doc_tipo[0]} docentes únicos, siendo la modalidad de mayor alcance por su formato "
    f"breve y accesible. El Diplomado, con {n_doc_tipo[1]} participantes únicos, implica mayor "
    f"compromiso temporal y profundidad en el desarrollo de competencias pedagógicas.")
bajada(doc,
    f"El Proyecto agrupa a {n_doc_tipo[2]} docentes únicos en {n_inst_tipo[2]} instancias, "
    f"con algunos participando en más de un proyecto. Esta modalidad, gestionada por VRIIP, "
    f"apunta al desarrollo de investigación e innovación en docencia, representando la "
    f"dimensión más especializada del Plan de Formación.")
doc.add_page_break()

# 3.2
titulo(doc, "3.2  Tipo de Formación por Jerarquía Académica")
intro(doc,
    "El tipo de instancia de perfeccionamiento varía según la jerarquía del docente: "
    "los rangos superiores muestran mayor diversificación entre modalidades, mientras "
    "los rangos de entrada concentran su participación principalmente en Talleres. "
    "Esta distribución refleja tanto la disponibilidad de tiempo como el perfil "
    "formativo esperado en cada etapa de la carrera académica.")
casc = "357 docentes formados · Tipo de formación por jerarquía\n"
for j in jer_validos2:
    nt = part[part["jerarquia"]==j]["rut_key"].nunique()
    ntal = part[(part["tipo_formacion"]=="TALLER")&(part["jerarquia"]==j)]["rut_key"].nunique()
    ndip = part[(part["tipo_formacion"]=="DIPLOMADO")&(part["jerarquia"]==j)]["rut_key"].nunique()
    npro = part[(part["tipo_formacion"]=="PROYECTO")&(part["jerarquia"]==j)]["rut_key"].nunique()
    casc += f"  ├── {ABREV_JER[j]:18}: n={nt:3d} | T={ntal} D={ndip} P={npro}\n"
cascada(doc, casc)
grafico(doc, pngs["3.2"])
bajada(doc,
    "Los rangos de entrada al escalafón (Instructor Docente y Asistente Docente) concentran "
    "su participación principalmente en Talleres, la modalidad más accesible y de menor "
    "compromiso temporal. Esto es coherente con la etapa de adaptación institucional "
    "que caracteriza a estos docentes.")
bajada(doc,
    "Los rangos superiores (Asociado y Titular) muestran mayor participación relativa en "
    "Diplomados y Proyectos, modalidades que requieren mayor dedicación y están más "
    "alineadas con el perfil investigativo y de liderazgo académico esperado en los "
    "niveles avanzados del escalafón.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 4
# ══════════════════════════════════════════════════════════════════════════════
titulo(doc, "SECCIÓN 4 — Brechas de Acceso al Perfeccionamiento")

# 4.1
titulo(doc, "4.1  Perfil de los Docentes que No Han Participado")
intro(doc,
    f"Los {n_no_part} docentes que no han participado en ninguna instancia de "
    f"perfeccionamiento entre 2022 y 2025 ({100*n_no_part/n_total:.1f}% del universo) "
    f"constituyen el principal foco de atención para la política institucional de "
    f"desarrollo académico. Su perfil por jerarquía y antigüedad permite identificar "
    f"grupos prioritarios para estrategias de convocatoria diferenciadas.")
casc  = f"917 docentes · {n_no_part} sin ninguna instancia registrada\n"
casc += f"  Por jerarquía (top 3 en términos absolutos):\n"
no_jer_sorted = sorted(zip(jer_validos,n_jer_no2),key=lambda x:-x[1])[:3]
for j,n in no_jer_sorted:
    casc += f"    ├── {ABREV_JER[j]:18}: {n} sin participación\n"
casc += f"  Por antigüedad:\n"
for a,n in zip(ORD_ANT,n_ant_no2):
    casc += f"    ├── {a:5} años: {n} sin participación\n"
cascada(doc, casc)
grafico(doc, pngs["4.1"])
bajada(doc,
    f"En términos absolutos, los grupos con mayor número de docentes sin participación "
    f"son {ABREV_JER[no_jer_sorted[0][0]]} ({no_jer_sorted[0][1]} docentes), "
    f"{ABREV_JER[no_jer_sorted[1][0]]} ({no_jer_sorted[1][1]}) y "
    f"{ABREV_JER[no_jer_sorted[2][0]]} ({no_jer_sorted[2][1]}). Dado que Instructor "
    f"Docente es el grupo más numeroso del universo, concentra naturalmente el mayor "
    f"volumen de no participantes.")
bajada(doc,
    f"Los docentes de 0-4 años de antigüedad registran {n_ant_no2[0]} sin participación, "
    f"el mayor volumen absoluto por tramo. Este grupo, en etapa de incorporación "
    f"institucional, podría beneficiarse de programas de inducción formativa obligatoria "
    f"que garanticen una base mínima de perfeccionamiento desde el inicio de la carrera.")
doc.add_page_break()

# 4.2
titulo(doc, "4.2  Intensidad de Participación")
intro(doc,
    f"Entre los {n_part} docentes que han participado en al menos una instancia, "
    f"{vals_int[0]} ({pct_int[0]:.0f}%) lo ha hecho en una sola oportunidad, "
    f"evidenciando que el perfeccionamiento es principalmente un evento puntual "
    f"y no una práctica sostenida. Solo el {100-acum[1]:.0f}% ha participado en "
    f"3 o más instancias, configurando un grupo reducido de docentes con alta "
    f"dedicación formativa.")
cascada(doc,
    f"{n_part} docentes con al menos 1 instancia · Distribución por frecuencia\n"
    f"  ├── 1 instancia  : {vals_int[0]:3d} ({pct_int[0]:.0f}%)\n"
    f"  ├── 2 instancias : {vals_int[1]:3d} ({pct_int[1]:.0f}%)\n"
    f"  ├── 3 instancias : {vals_int[2]:3d} ({pct_int[2]:.0f}%)\n"
    f"  └── 4 o más      : {vals_int[3]:3d} ({pct_int[3]:.0f}%)")
grafico(doc, pngs["4.2"])
bajada(doc,
    f"El {pct_int[0]:.0f}% de los docentes formados ({vals_int[0]}) ha participado en una "
    f"sola instancia, lo que sugiere que el perfeccionamiento es percibido principalmente "
    f"como un evento aislado. Fomentar la reincidencia — a través de rutas formativas "
    f"progresivas o compromisos anuales de desarrollo — permitiría profundizar el impacto "
    f"del Plan de Formación.")
bajada(doc,
    f"El {pct_int[3]:.0f}% ({vals_int[3]} docentes) con 4 o más instancias representa "
    f"el núcleo de alta dedicación formativa de UCEN. Identificar y visibilizar este "
    f"grupo como referentes de desarrollo docente podría contribuir a generar una cultura "
    f"de perfeccionamiento continuo en el cuerpo académico.")

# ── Guardar ───────────────────────────────────────────────────────────────────
doc.save(OUT_DOC)
print(f"\nEntregable P2 guardado: {OUT_DOC}")
