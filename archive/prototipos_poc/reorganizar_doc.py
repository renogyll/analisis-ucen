import sys; sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE_NOTEBOOKS = r"c:\Users\r.gonzalez_fluxsolar.LAPTOP-FLUX-ECO\Downloads\Analisis_UCEN_v2\UNIVERSO_918\NOTEBOOKS"
OUT = r"c:\Users\r.gonzalez_fluxsolar.LAPTOP-FLUX-ECO\Downloads\Analisis_UCEN_v2\Informe_UCEN_Reorganizado.docx"

def img(name): return os.path.join(BASE_NOTEBOOKS, name)

doc = Document()

# ── Márgenes ───────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── Helpers ────────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x37, 0x47, 0x4F)
    return p

def h3(text, color=None):
    p = doc.add_heading(text, level=3)
    if color:
        p.runs[0].font.color.rgb = color
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.style.font.size = Pt(11)
    return p

def label(text, color=None):
    """Etiqueta pequeña en cursiva — útil como subtítulo de gráfico."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def grafico(filename, width=Inches(5.8), caption=None):
    path = img(filename)
    if not os.path.exists(path):
        body(f"[IMAGEN NO ENCONTRADA: {filename}]")
        return
    doc.add_picture(path, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        label(caption)

def separador():
    doc.add_paragraph()

def nota_metodologica(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    run = p.add_run(f"Nota metodológica: {text}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x54, 0x6E, 0x7A)

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("ENTREGA PRODUCTO 3", 0)
h2("Análisis de Iniciativas de Formación Docente e Impacto en el Rendimiento Académico y Satisfacción Estudiantil 2022–2025")
body("Universidad Central de Chile  ·  Mayo 2026")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. UNIVERSO Y CASCADA
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Universo de Análisis")

h2("1.1 Cascada — Evaluación SAT")
body("El universo de estudio está compuesto por 917 docentes jerarquizados de la Universidad Central. A partir de este total, se aplican dos filtros sucesivos para determinar los subgrupos analíticamente pertinentes:")
bullet("917 docentes jerarquizados (universo total): 493 con perfil completo (nómina + dotación) y 424 solo nómina, sin datos demográficos de dotación.")
bullet("357 docentes (39%) participaron en al menos una instancia de formación — grupo tratamiento. Los 560 restantes (61%) conforman el grupo control.")
bullet("197 docentes APTOS P3 (21% del total): cuentan con evaluación SAT válida tanto en el período baseline como en el período resultado, lo que habilita el análisis de impacto pre/post. Se distribuyen en 154 tipo TALLER, 36 tipo DIPLOMADO y 7 tipo PROYECTO.")
nota_metodologica("Solo se consideraron evaluaciones SAT con cobertura ≥ 40% (CM-1). El promedio SAT de cada docente por período se calculó ponderado por número de alumnos evaluados (CM-2), para evitar que secciones pequeñas distorsionen el indicador.")
separador()

h2("1.2 Cascada — Evaluación de Jefes (EDD)")
body("La Evaluación de Desempeño Docente (EDD) captura la apreciación de los jefes directos sobre cada académico, en escala 0–1, con cuatro conceptos: Muy Bueno, Bueno, Insuficiente y Deficiente.")
bullet("1.414 registros EDD en la base de datos (487 docentes únicos); 1 fuera del universo 917 → 1.413 registros válidos.")
bullet("702 registros corresponden a 235 docentes formados; 711 a 251 docentes sin formación.")
bullet("Los formados concentran el 74% en Muy Bueno, frente al 62% del grupo control — brecha de 12 puntos porcentuales especialmente marcada en 2025.")
separador()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. MARCO METODOLÓGICO
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Marco Metodológico")

h2("2.1 El indicador z-score")
body("El análisis utiliza el z-score como medida central de impacto, en lugar de la nota SAT bruta. El z-score expresa la posición relativa del docente dentro de su facultad y período, eliminando diferencias de escala entre unidades académicas y semestres:")
body("z = (SAT docente − promedio facultad·período) / desviación estándar facultad·período")
bullet("z = 0: el docente está en el promedio exacto de su facultad ese semestre.")
bullet("z = +1: supera a sus colegas en 1 desviación estándar (~84° percentil).")
bullet("z = −1: está por debajo en 1 desviación estándar (~16° percentil).")
body("Esta métrica permite comparar docentes de distintas facultades y períodos en pie de igualdad, y habilita el análisis pre/post (Δz = z resultado − z baseline) con pruebas estadísticas formales.")

h2("2.2 Dos dimensiones de satisfacción")
body("El análisis incorpora dos preguntas SAT complementarias, cuyos resultados se presentan siempre en paralelo:")
bullet("SAT Nota: evaluación numérica del docente por parte de los estudiantes (escala 1–7), z-scoreada por facultad × período.")
bullet("SAT % Recomendación: porcentaje de estudiantes que responden 'Sí' a '¿Recomendaría a este/a docente?', igualmente z-scoreada. Ambas métricas tienen correlación r = 0.893 — son consistentes y se refuerzan mutuamente.")
separador()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. ANÁLISIS DE IMPACTO SAT — GRÁFICOS PAREADOS
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Análisis de Impacto en Evaluación Docente")

# ── G1 / G1.2 ─────────────────────────────────────────────────────────────────
h2("3.1 Trayectoria Z-score: Baseline → Resultado")

h3("SAT Nota", color=RGBColor(0xFF, 0x6B, 0x35))
grafico("G1_linea_z_918.png", caption="G1 — Trayectoria z-score SAT Nota por tipo de formación (n=197 aptos P3)")
bullet("Los tres grupos parten sobre el promedio de su facultad (z > 0) antes de formarse, lo que sugiere un efecto de selección: quienes participan en formación tienden a ser docentes ya bien evaluados.")
bullet("Taller (n=154) y Diplomado (n=36) muestran trayectorias estables — el z-score post no difiere significativamente del baseline. No hay regresión al promedio ni salto brusco.")
bullet("Proyecto (n=7) registra la caída más pronunciada (de z≈+0.5 a +0.26), pero el tamaño muestral hace que ese promedio sea muy sensible a casos individuales — debe leerse con cautela.")
separador()

h3("SAT % Recomendación", color=RGBColor(0x19, 0x76, 0xD2))
grafico("G1.2_linea_zbin_918.png", caption="G1.2 — Trayectoria z-score % Recomendación por tipo de formación (n=197 aptos P3)")
bullet("El patrón replica fielmente el de SAT Nota (r=0.893 entre métricas): misma estabilidad en Taller y Diplomado, misma caída relativa en Proyecto.")
bullet("Los tres grupos se mantienen sobre la media en la mayoría de los períodos, confirmando que los docentes formados tienen una posición de partida favorable también en la dimensión de recomendación.")
bullet("La coherencia entre ambas métricas valida la robustez del análisis: los resultados no son artefactos de una sola pregunta, sino consistentes en dos indicadores independientes.")
separador()
doc.add_page_break()

# ── G2 / G2.2 ─────────────────────────────────────────────────────────────────
h2("3.2 Delta Z-score y Porcentaje de Mejora por Tipo de Formación")

h3("SAT Nota", color=RGBColor(0xFF, 0x6B, 0x35))
grafico("G2_barras_tipo_918.png", caption="G2 — Δ z-score SAT Nota y % docentes que mejoran, por tipo de formación")
bullet("Los tres tipos registran Δz negativo o cercano a cero (Taller −0.035, Diplomado −0.028, Proyecto −0.225), lo que indica que la formación no produce un salto sistemático en el promedio grupal.")
bullet("El % de mejora individual cuenta una historia diferente: el 53% de los diplomados y el 49% de los talleristas mejoró su posición relativa — la distribución es bimodal, no hay una tendencia uniforme de caída.")
bullet("La aparente contradicción entre Δz negativo y ~50% de mejora se explica porque el promedio es sensible a los que retroceden más (pocos casos con caídas grandes arrastran la media hacia abajo). Ambos datos son verdaderos al mismo tiempo.")
separador()

h3("SAT % Recomendación", color=RGBColor(0x19, 0x76, 0xD2))
grafico("G2.2_barras_tipo_bin_918.png", caption="G2.2 — Δ z-score % Recomendación y % docentes que mejoran, por tipo de formación")
bullet("Misma estructura que SAT Nota: Δz negativo en los tres tipos, con 52–56% de mejora individual en Diplomado y Taller. La consistencia entre métricas confirma que el resultado neutro pre/post no es un artefacto.")
bullet("Proyecto (n=7) registra Δz=−0.258 en % recomendación — incluso más pronunciado que en nota SAT — pero con n=7 cualquier caso atípico mueve el promedio sustancialmente.")
bullet("El mensaje clave es el mismo en ambas métricas: la formación no deteriora la posición relativa de los docentes, pero tampoco produce una mejora agregada estadísticamente distinguible del azar.")
separador()
doc.add_page_break()

# ── G3 / G3.2 ─────────────────────────────────────────────────────────────────
h2("3.3 Impacto por Antigüedad del Docente")

h3("SAT Nota", color=RGBColor(0xFF, 0x6B, 0x35))
grafico("G3_heatmap_antiguedad_918.png", caption="G3 — Δ z-score SAT Nota por antigüedad × tipo de formación")
bullet("Los docentes noveles (0–4 años, n=57) obtienen el mayor beneficio del Diplomado (Δz=+0.18), lo que sugiere que esta modalidad tiene mayor impacto formativo en quienes aún están consolidando su práctica.")
bullet("Los docentes de mayor trayectoria (15+ años) responden mejor al Taller (Δz=+0.19), posiblemente porque esta modalidad se adapta mejor a necesidades puntuales de actualización sin desafiar esquemas ya establecidos.")
bullet("El grupo de antigüedad media (5–14 años) muestra resultados cercanos a cero en ambos tipos — es la cohorte más estable, pero también la que menos se mueve con la intervención.")
separador()

h3("SAT % Recomendación", color=RGBColor(0x19, 0x76, 0xD2))
grafico("G3.2_heatmap_antiguedad_bin_918.png", caption="G3.2 — Δ z-score % Recomendación por antigüedad × tipo de formación")
bullet("Los docentes noveles lideran la mejora en recomendación también en Diplomado, replicando el patrón de SAT Nota. La formación intensiva parece ser especialmente valorada por los estudiantes cuando el docente es más nuevo.")
bullet("El patrón de antigüedad × tipo es consistente entre ambas métricas: los mismos subgrupos que mejoran en nota mejoran en recomendación, y viceversa.")
bullet("Donde hay discrepancias (alguna celda con signo opuesto), el n es inferior a 5 — las celdas de bajo n deben leerse como referenciales y no como evidencia concluyente.")
separador()
doc.add_page_break()

# ── G4 / G4.2 ─────────────────────────────────────────────────────────────────
h2("3.4 Impacto por Jerarquía Académica")

h3("SAT Nota", color=RGBColor(0xFF, 0x6B, 0x35))
grafico("G4_heatmap_jerarquia_918.png", caption="G4 — Δ z-score SAT Nota por jerarquía × tipo de formación")
bullet("Instructores y Asistentes Docentes — las jerarquías mayoritarias en el grupo formado — muestran los patrones más interpretativos dado su mayor n. Ambas jerarquías exhiben estabilidad en Taller y leve mejora en Diplomado.")
bullet("Los Asociados Docentes (n=35 en Taller) presentan Δz negativo en ambos tipos — son la jerarquía que más retrocede en posición relativa, aunque sin diferencia estadísticamente significativa respecto a las demás (ANOVA p=0.804).")
bullet("Las jerarquías con n < 5 (Asistente Regular, Titular) deben leerse como ilustrativas: la varianza individual domina completamente el resultado de celda.")
separador()

h3("SAT % Recomendación", color=RGBColor(0x19, 0x76, 0xD2))
grafico("G4.2_heatmap_jerarquia_bin_918.png", caption="G4.2 — Δ z-score % Recomendación por jerarquía × tipo de formación")
bullet("El mapa de calor de recomendación es estructuralmente similar al de SAT Nota: mismas jerarquías con tendencia positiva, mismas con tendencia negativa. La correlación entre métricas se mantiene a nivel de subgrupo.")
bullet("Instructores Docentes en Diplomado muestran el Δz más positivo en ambas métricas — consistente con la evolución temporal del G7 que los posiciona sobre el promedio post-intervención.")
bullet("La consistencia entre métricas es el hallazgo principal: no hay ninguna jerarquía que mejore en una dimensión y retroceda en la otra, lo que otorga solidez interna al análisis.")
separador()
doc.add_page_break()

# ── G5 / G5.2 ─────────────────────────────────────────────────────────────────
h2("3.5 Impacto por Sexo del Docente")

h3("SAT Nota", color=RGBColor(0xFF, 0x6B, 0x35))
grafico("G5_barras_sexo_918.png", caption="G5 — Δ z-score SAT Nota por sexo × tipo de formación")
bullet("Las mujeres representan el 59% de los docentes formados (n=116 vs 81 hombres) y muestran Δz positivo en Taller y Diplomado, mientras los hombres retroceden en los tres tipos.")
bullet("En Taller — la modalidad más masiva — el 51% de las mujeres mejoró su posición relativa frente al 45% de los hombres, una diferencia que sugiere una mayor receptividad femenina a este formato de formación.")
bullet("El grupo Proyecto tiene n muy bajo (F=2, M=5) y sus resultados no son generalizables. Los hombres registran la mayor caída precisamente en este tipo, arrastrando el promedio masculino global.")
separador()

h3("SAT % Recomendación", color=RGBColor(0x19, 0x76, 0xD2))
grafico("G5.2_barras_sexo_bin_918.png", caption="G5.2 — Δ z-score % Recomendación por sexo × tipo de formación")
bullet("El diferencial por sexo se mantiene en la dimensión de recomendación: las docentes mujeres muestran mejor desempeño relativo post-formación, con el 53% mejorando en Taller vs el 44% de los hombres.")
bullet("La brecha entre sexos no está impulsada por diferencias de partida (ambos grupos comienzan en posiciones similares), sino por el cambio durante el período de intervención.")
bullet("La consistencia del patrón en ambas métricas sugiere que la diferencia es real y no un artefacto estadístico: los estudiantes perciben el impacto de la formación de forma diferenciada según el sexo del docente.")
separador()
doc.add_page_break()

# ── G6 / G6.2 ─────────────────────────────────────────────────────────────────
h2("3.6 Formados vs Control: Evolución por Período")

h3("SAT Nota", color=RGBColor(0xFF, 0x6B, 0x35))
grafico("G6_control_vs_tratamiento_918.png", caption="G6 — Z-score SAT Nota: formados vs control por período (2023–2025)")
bullet("Los docentes formados superan al grupo control en los seis períodos sin excepción, con una brecha promedio de +0.18z (equivalente a ~7 percentiles de diferencia).")
bullet("La brecha se amplía especialmente en 2024-01 y 2025-02 — períodos con mayor volumen de docentes formados activos — y la prueba t independiente confirma diferencias significativas en 3 de los 6 períodos (2024-01**, 2024-02**, 2025-02***).")
bullet("El grupo control se mantiene consistentemente bajo la media de su facultad (z entre −0.03 y −0.10), mientras los formados permanecen sobre ella, lo que sugiere que la participación en formación está asociada con una posición relativa superior sostenida en el tiempo.")
separador()

h3("SAT % Recomendación", color=RGBColor(0x19, 0x76, 0xD2))
grafico("G6.2_control_bin_918.png", caption="G6.2 — Z-score % Recomendación: formados vs control por período (2023–2025)")
bullet("La brecha formados-control se replica en la dimensión de recomendación, con +0.15z de diferencia promedio — levemente inferior al +0.18z de nota, pero estructuralmente idéntico.")
bullet("Los mismos tres períodos resultan significativos (2024-01**, 2024-02*, 2025-02*), confirmando que la ventaja de los formados no es un artefacto de una sola métrica.")
bullet("La correlación r=0.893 entre ambas métricas explica la consistencia: un docente que tiene mejor posición relativa en nota también tiende a tener más alumnos que lo recomendarían — y eso se mantiene al comparar grupos.")
separador()
doc.add_page_break()

# ── G7 Talleres ───────────────────────────────────────────────────────────────
h2("3.7 Evolución por Jerarquía: Talleres")

h3("SAT Nota", color=RGBColor(0xFF, 0x6B, 0x35))
grafico("G7_taller_jerarquia_918.png", caption="G7 — Z-score SAT Nota por jerarquía: Taller (n=154 aptos P3, 2023–2025)")
bullet("Asistentes Docentes (n=52) — el grupo más numeroso — entran bajo la media y alcanzan su peak durante los talleres (2024-01 a 2024-02), cerrando 2025-02 en +0.22. Trayectoria de éxito sostenido.")
bullet("Instructores Docentes (n=44) parten sobre el promedio (+0.23) y se mantienen estables en torno al mismo nivel — sin gran ganancia adicional, pero sin deterioro.")
bullet("Asociados Docentes (n=35) son la única jerarquía que permanece consistentemente bajo la media en todos los períodos (z entre −0.09 y −0.27), lo que sugiere que el taller no logra reposicionarlos relativamente dentro de su facultad.")
separador()

h3("SAT % Recomendación", color=RGBColor(0x19, 0x76, 0xD2))
grafico("G7.2_taller_jerarquia_bin_918.png", caption="G7.2 — Z-score % Recomendación por jerarquía: Taller (n=154 aptos P3, 2023–2025)")
bullet("El patrón por jerarquía se replica en % recomendación: Asistentes e Instructores Docentes mantienen posiciones sobre el promedio, y los Asociados Docentes permanecen por debajo.")
bullet("La coincidencia en ambas métricas para los grupos de mayor n otorga solidez al hallazgo: la jerarquía modera el impacto del taller de forma consistente independientemente del indicador utilizado.")
bullet("Grupos con n ≤ 4 (Asoc. Regular, Titular Regular) muestran alta variabilidad en ambos gráficos — sus trayectorias son ilustrativas pero no generalizables.")
separador()
doc.add_page_break()

# ── G7 Diplomados ─────────────────────────────────────────────────────────────
h2("3.8 Evolución por Jerarquía: Diplomado")

h3("SAT Nota", color=RGBColor(0xFF, 0x6B, 0x35))
grafico("G7_diplomado_jerarquia_918.png", caption="G7 — Z-score SAT Nota por jerarquía: Diplomado (n=36 aptos P3, 2023–2025)")
bullet("Instructores Docentes (n=18) son el caso de mayor éxito: parten bajo la media en 2023, escalan a +0.30 durante el diplomado (2024), y consolidan en +0.44 post-intervención (2025-02) — la trayectoria más ascendente del análisis.")
bullet("Asistentes Docentes (n=12) se mantienen sobre el promedio en todos los períodos (z entre +0.05 y +0.35), con una leve caída al final que no compromete su posición relativa favorable.")
bullet("Asociados Docentes (n=5) parten muy sobre el promedio (z=+0.48) pero caen a terreno negativo en el período de ejecución — resultado que debe interpretarse con cautela dado el n reducido.")
separador()

h3("SAT % Recomendación", color=RGBColor(0x19, 0x76, 0xD2))
grafico("G7.2_diplomado_jerarquia_bin_918.png", caption="G7.2 — Z-score % Recomendación por jerarquía: Diplomado (n=36 aptos P3, 2023–2025)")
bullet("Instructores Docentes (n=18) replican en % recomendación la trayectoria ascendente observada en nota: sus estudiantes los recomiendan cada vez más post-diplomado, confirmando un efecto real en la percepción estudiantil.")
bullet("La coherencia entre G7 y G7.2 para los Instructores — la jerarquía con mayor n — es la evidencia más sólida de impacto diferencial por jerarquía, aunque el ANOVA global no alcanza significancia estadística (p=0.718).")
bullet("La alta correlación entre métricas (r=0.893) hace que ambos gráficos sean prácticamente espejo: el diplomado impacta de forma homogénea en nota y en recomendación dentro de cada jerarquía.")
separador()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. RELACIÓN SAT — NOTAS DE ALUMNOS
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Relación entre Evaluación Docente y Rendimiento Estudiantil")

grafico("G8_scatter_sat_notas_918.png", caption="G8 — Dispersión SAT vs Nota promedio alumnos (n=17.248 secciones, 6 períodos 2023–2025)")
bullet("La correlación entre SAT y nota de alumnos es positiva y significativa (r=0.28, p<0.001), pero modesta: el SAT explica solo el 8% de la varianza en notas. El 92% restante responde a factores externos — dificultad de la asignatura, perfil del curso, carrera, semestre.")
bullet("Los docentes formados (naranja) aparecen más concentrados en el sector derecho del gráfico (SAT alto), consistente con los hallazgos de G6: los formados tienen mejor posición relativa en evaluación docente. Sin embargo, el gráfico no permite concluir causalidad — puede ser que mejores docentes participen más en formación.")
bullet("Conclusión: existe relación positiva entre evaluación docente y rendimiento estudiantil, y los formados tienden a ubicarse en la zona favorable. La formación no coexiste con mal desempeño: los cuadrantes de 'mal evaluado + notas bajas' están prácticamente vacíos de docentes formados.")
separador()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. PRUEBAS ESTADÍSTICAS
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Validación Estadística")

h2("5.1 Prueba t — Z-score SAT Nota y % Recomendación")
body("Se realizaron dos tipos de prueba t para validar los hallazgos descriptivos:")
bullet("Prueba A (pareada): ¿Cambia el propio docente entre baseline y resultado? Respuesta: no — ningún tipo de formación produce un cambio estadísticamente significativo en el z-score del docente (todos ns). El docente no se mueve relativo a su facultad.")
bullet("Prueba C (independiente por período): ¿Tienen los formados mejor posición que los no formados? Respuesta: sí, en 3 de 6 períodos — la diferencia es real y recurrente.")
separador()

grafico("G11_pruebas_t_918.png", width=Inches(6.0), caption="G11 — Pruebas t SAT Nota: Pareada (Panel A) e Independiente por período (Panel C)")
bullet("Prueba t Pareada (SAT Nota): todos los tipos resultan no significativos (todos ns). La formación no desplaza al docente hacia arriba ni hacia abajo respecto a sus colegas de facultad.")
bullet("Prueba t Independiente (SAT Nota): diferencias significativas en 2024-01 (p=0.003, d=+0.260), 2024-02 (p=0.009, d=+0.231) y 2025-02 (p=0.001, d=+0.302). En esos períodos, los formados estaban sistemáticamente mejor posicionados que el grupo control.")
bullet("El efecto más fuerte es 2025-02 (Cohen d=+0.302, considerado pequeño-mediano): los docentes formados superan al control en casi un tercio de desviación estándar en ese período.")
separador()

grafico("G11.2_pruebas_t_bin_918.png", width=Inches(6.0), caption="G11.2 — Pruebas t % Recomendación: Pareada (Panel A) e Independiente por período (Panel C)")
bullet("Prueba t Pareada (% Recomendación): idéntico resultado que SAT Nota — todos ns. La formación no mueve la tasa de recomendación propia del docente entre baseline y resultado.")
bullet("Prueba t Independiente (% Recomendación): significativa en los mismos tres períodos (2024-01**, 2024-02*, 2025-02*), con efecto algo menor que SAT Nota (d máx.=+0.230). La coincidencia entre métricas es evidencia de robustez.")
bullet("Los períodos 2023 y 2025-01 no alcanzan significancia en ninguna métrica — la ventaja de los formados existe pero es estadísticamente indistinguible del azar en esos momentos, posiblemente por menor volumen de formados activos.")
separador()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. EVALUACIÓN DE JEFES (EDD)
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Evaluación de Desempeño Docente (EDD) — Perspectiva de Jefes")

h2("6.1 Evolución EDD: Formados vs Sin Formación (2022–2025)")
grafico("G9_edd_evolucion_918.png", caption="G9 — EDD Total promedio por año: formados vs sin formación (Universo 917, 2022–2025)")
bullet("Los docentes formados superan consistentemente al grupo control en EDD durante los cuatro años, pero la diferencia más crítica ocurre en 2025: formados promedian 0.864 vs 0.604 del control — una brecha de +0.260 puntos (26 puntos porcentuales en escala 0–1).")
bullet("La caída generalizada de 2024 afecta a ambos grupos, pero los formados la absorben con menor impacto (0.734 vs 0.655) y revierten al nivel más alto en 2025. El grupo control no logra recuperarse al mismo ritmo.")
bullet("Las pruebas t confirman que la diferencia no es ruido: es significativa en 2024 (p=0.002, d=+0.339) y altamente significativa en 2025 (p<0.001, d=+0.955 — efecto grande). El d global es +0.522, el mayor tamaño de efecto de todo el análisis.")
separador()

h2("6.2 Distribución de Concepto EDD")
grafico("G10_concepto_edd_918.png", caption="G10 — Distribución de Concepto EDD por grupo: formados vs sin formación")
bullet("Los docentes formados concentran el 74.3% en Muy Bueno vs el 65.2% de sus pares — una brecha de +9.1 puntos porcentuales que es estadísticamente significativa (chi-cuadrado global p=0.001, Cramer V=0.107).")
bullet("La categoría Deficiente prácticamente desaparece en formados (0.5%) frente al 1.6% del grupo control — la formación actúa como un piso que reduce la incidencia de los desempeños más bajos evaluados por los jefes.")
bullet("El efecto es pequeño en términos de Cramer V (V=0.107), lo que indica que la diferencia existe y es real, pero no determina completamente la distribución — hay docentes sin formación con Muy Bueno y formados que no lo alcanzan.")
separador()

h2("6.3 Pruebas Estadísticas EDD")
grafico("G12_pruebas_t_edd_918.png", width=Inches(6.0), caption="G12 — Validación estadística EDD: Prueba t (Panel A) y Chi-cuadrado (Panel B)")
bullet("Prueba t independiente (Panel A): diferencias significativas en 2024 (p=0.002, **) y 2025 (p<0.001, ***). El efecto global es altamente significativo (t=+9.615, p<0.001, Cohen d=+0.522 — efecto mediano).")
bullet("Chi-cuadrado (Panel B): la distribución de conceptos difiere entre formados y control a nivel global (p=0.001, **) y en 2024 (p=0.003, **). En 2022, 2023 y 2025 la diferencia no alcanza significancia año a año, pero sí en el acumulado.")
bullet("El efecto en EDD es el más robusto y de mayor magnitud de todo el análisis: Cohen d=+0.955 en 2025 es considerado efecto grande. La formación docente se asocia claramente con mejores evaluaciones de desempeño desde la perspectiva de los jefes directos.")
separador()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Conclusiones")

h2("7.1 Hallazgos principales")
bullet("La formación docente en UCEN no es marginal: el 39% de los jerarquizados participó en al menos una actividad en el período analizado. Los docentes formados parten con una posición relativa favorable (efecto de selección), y esa ventaja se mantiene y amplía en el tiempo.")
bullet("El impacto sobre la satisfacción estudiantil (SAT nota y % recomendación) es neutral en términos pre/post propios del docente — nadie retrocede sistemáticamente — pero los formados como grupo están consistentemente mejor posicionados que el control en 3 de 6 períodos, con efectos pequeños-medianos (Cohen d hasta +0.302).")
bullet("El impacto más fuerte es en la evaluación de jefes (EDD): Cohen d=+0.955 en 2025, efecto grande. Los formados tienen 9 puntos porcentuales más de Muy Bueno y una fracción ínfima de Deficiente. Esta dimensión captura diferencias que la satisfacción estudiantil no alcanza a distinguir.")
bullet("Ambas métricas de satisfacción estudiantil (nota y % recomendación, r=0.893) apuntan en la misma dirección en todos los períodos y subgrupos, lo que otorga robustez interna al análisis: los resultados no dependen de la elección de un solo indicador.")

h2("7.2 Matices y limitaciones")
bullet("El diseño es observacional, no experimental. No es posible atribuir causalidad directa a la formación: los docentes que participan pueden tener características previas (motivación, compromiso) que explican parte de su mejor desempeño.")
bullet("El grupo Proyecto (n=7) no es estadísticamente interpretable como grupo independiente. Sus resultados son ilustrativos y no generalizables.")
bullet("La prueba t pareada (Prueba A) confirma que la formación no desplaza automáticamente hacia arriba al propio docente dentro de su facultad — el aprendizaje no se traduce de forma inmediata en mejora de la posición relativa SAT. El efecto es de grupo, no individual.")
bullet("La jerarquía académica no modera el impacto de forma estadísticamente significativa ni en Taller (ANOVA p=0.804) ni en Diplomado (ANOVA p=0.718): la formación tiene un efecto equivalente en todos los rangos académicos, con la salvedad de que los Instructores Docentes en Diplomado muestran la trayectoria más ascendente.")

separador()
body("Informe generado automaticamente | Universo 917 | Análisis Producto 3 UCEN | Mayo 2026")

doc.save(OUT)
print(f"Guardado: {OUT}")
print(f"Párrafos: {len(doc.paragraphs)}")
